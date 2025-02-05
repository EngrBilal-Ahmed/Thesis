from docx import Document

# Create a Word document
doc = Document()

# Add a title to the document
doc.add_heading('Symbolic Representation of Algorithm', level=1)

# Add a table with clear formatting
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'

# Add headers
headers = ["Phase", "Patient (Smart Card)", "Trusted Server (TS)", "Medical Server (MS)"]
for idx, header in enumerate(headers):
    table.cell(0, idx).text = header

# Fill in the rows
rows = [
    [
        "Registration Phase",
        "- Selects IDi, pwi, and Bi (biometric data).\n"
        "- Computes:\n"
        "  HPW = h(IDi ⊕ pwi ⊕ h_bio(Bi))\n"
        "  Bir = Bi ⊕ r1\n"
        "  P = HPW ⊕ Bir\n"
        "  R = h(IDi || h_bio(Bi) || pwi) ⊕ r1\n"
        "- Sends IDi, HPW, h(·), P to TS.",
        "- Computes:\n"
        "  M1 = IDi ⊕ IDt\n"
        "  Cut = E_k1(M1 || r2 || P)\n"
        "  Ni = h(IDi || IDt || k1 || r2) ⊕ HPW\n"
        "- Sends Cut, Ni to Ui.",
        ""
    ],
    [
        "Login Phase",
        "- Inputs IDi, pwi, Bi*.\n"
        "- Computes r1 = R ⊕ h(IDi || h_bio(Bi) || pwi).\n"
        "- Verifies Bi using Δ(Bi, Bi*).\n"
        "- If valid, computes:\n"
        "  HPW' = h(IDi ⊕ pwi ⊕ h_bio(Bi*))\n"
        "  P' = HPW' ⊕ Bir\n"
        "  Ai = h(IDi || (m ⊕ IDm)) ⊕ P'\n"
        "  X1 = h(IDi || Cut || Ai) ⊕ m\n"
        "- Sends Cut, Authut = E_SKut(X1 || IDm || Ai) to TS.",
        "",
        ""
    ],
    [
        "Authentication Phase",
        "",
        "- Verifies Authut.\n"
        "- Extracts X1, IDm, Ai.\n"
        "- Computes m' = X1 ⊕ h(IDi || Cut || Ai).\n"
        "- If valid, computes SKut = h(IDi || IDt || k1 || r2).\n"
        "- Sends Authtm = E_SKtm(M2 || X2) to MS.",
        "- Verifies Authtm.\n"
        "- Computes SKtm = h(IDm || n || (X2 ⊕ IDm)).\n"
        "- Updates SKtm and sends Authtu to TS.\n"
        "- Session key SK = h(IDm || n || (IDi ⊕ Ai ⊕ m)) is shared."
    ]
]

# Populate the table
for row_idx, row in enumerate(rows, start=1):
    for col_idx, cell in enumerate(row):
        table.cell(row_idx, col_idx).text = cell

# Save the document to a file
file_path = "D:/Symbolic_Representation_Algorithm_Readable.docx"
doc.save(file_path)

file_path
