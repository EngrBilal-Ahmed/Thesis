from docx import Document

# Create a Word document
doc = Document()

# Add a title for the nomenclature
doc.add_heading('Nomenclature', level=1)

# Add a table for the symbols and descriptions
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Add headers to the table
headers = ["Symbol", "Description"]
for idx, header in enumerate(headers):
    table.cell(0, idx).text = header

# Populate the table with symbols and descriptions
symbols = [
    ("IDi", "Patient Identity"),
    ("pwi", "Patient Password"),
    ("Bi", "Biometric Data"),
    ("r1", "Random Number 1"),
    ("r2", "Random Number 2"),
    ("M1", "Message 1"),
    ("Cut", "Encrypted Message"),
    ("Ni", "Hashed Key"),
    ("HPW", "Hashed Password"),
    ("P", "Session Parameter"),
    ("R", "Registration Parameter"),
    ("Ai", "Authentication Parameter"),
    ("X1", "Intermediate Value 1"),
    ("Authut", "Authentication Token to TS"),
    ("SKut", "Session Key (TS & User)"),
    ("Authtm", "Authentication Token to MS"),
    ("SKtm", "Session Key (TS & MS)"),
    ("SK", "Final Session Key"),
    ("h(.)", "Hash Function"),
    ("Ek(.)", "Encryption Function"),
    ("⊕", "XOR Operation"),
    ("||", "Concatenation Operation"),
    ("Δ", "Difference Function"),
]

# Fill the table with symbols and descriptions
for symbol, description in symbols:
    row = table.add_row().cells
    row[0].text = symbol
    row[1].text = description

# Save the document to a file
file_path = "D:/Nomenclature_Readable.docx"
doc.save(file_path)

file_path
